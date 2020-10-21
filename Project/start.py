import docx
template = "template.docx"


class CommonTempl():
	def __init__(self, template):	
		self.tempObj = docx.Document(template)
		self.getRawDictionary()

	def getRawDictionary(self):
		rows = self.tempObj.tables[0].rows
		dic = {}
		for i, a in enumerate(rows):
			dic[self.tempObj.tables[0].cell(i,1).text] = self.tempObj.tables[0].cell(i,3).text
		self.dictTempl = dic

class BelAgroProm(CommonTempl):
	def __init__(self, template, rawTemp, savePath):
		super(BelAgroProm, self).__init__(template)
		self.rawTemp = rawTemp
		self.savePath = savePath

		self.array = [
		self.dictTempl['fullName'], #1
		self.dictTempl['regNum']+", "+self.dictTempl['regDate']+", "+self.dictTempl['nameOfRegAuth'], #2
		self.dictTempl['legAddress'].split(',')[0], #3
		"Нет", #4
		self.dictTempl['regNum'], #5
		self.dictTempl['legAddress'], #6
		"-", #7
		"-", #8
		self.dictTempl['director']+"\n"+self.dictTempl['accountant'], #9
		self.dictTempl['manageStructure'], #10
		self.dictTempl['sizeOfAuthCap'], #11
		self.dictTempl['z'], #12
		self.dictTempl['z'], #13
		self.dictTempl['z'], #14
		self.dictTempl['z'], #15
		self.dictTempl['z'], #16
		self.dictTempl['z'], #17
		self.dictTempl['z'], #18
		self.dictTempl['z'], #19
		self.dictTempl['z'], #20
		self.dictTempl['z'], #21
		self.dictTempl['z'], #22
		self.dictTempl['z'], #23
		self.dictTempl['z'], #24
		self.dictTempl['z'], #25
		self.dictTempl['z'], #26
		self.dictTempl['z'], #27
		self.dictTempl['z'], #28
		self.dictTempl['z'], #29
		self.dictTempl['z'], #30
		self.dictTempl['z'], #31
		self.dictTempl['z'], #32
		self.dictTempl['z'],] #33

	def insertData(self):
		rawTemp = docx.Document(self.rawTemp)
		rows = rawTemp.tables[0].rows
		for i, a in enumerate(rows):
			
			rawTemp.tables[0].cell(i,1).text = self.array[i]
		rawTemp.save(self.savePath)

if __name__ == "__main__":
	a = BelAgroProm(template, 'rawTemplates/belagr.docx', 'complete/belagr.docx')
	a.insertData()
